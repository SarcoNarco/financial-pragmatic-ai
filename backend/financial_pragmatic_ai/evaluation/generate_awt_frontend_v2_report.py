"""Generate an Advanced Web Technology report focused on frontend_v2."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


TEMPLATE_PATH = Path("/Users/saroshnadaf/Downloads/Paper_Format.docx")
OUTPUT_PATH = Path(__file__).resolve().parent / "awt_frontend_v2_report.docx"


TITLE = "Advanced Web Technology Project Report: Frontend V2 of Financial Pragmatic AI"
AUTHOR = "Prepared for Advanced Web Technology"
AFFILIATION_LINES = [
    "Project: Financial Pragmatic AI Frontend V2",
    "Focus Area: Modern Frontend Design, Authentication, API Integration, and Interactive Dashboard Engineering",
]

ABSTRACT = (
    "This report presents the frontend_v2 implementation of the Financial Pragmatic AI project from the "
    "perspective of Advanced Web Technology. The V2 frontend is a modern React-based analytical dashboard built "
    "with Vite, Tailwind CSS, Supabase authentication, Framer Motion, Lucide icons, and Recharts-based data "
    "visualization. Its purpose is to provide a responsive, interactive user interface for financial transcript "
    "analysis, allowing users to sign in, submit transcript text, upload local transcript files, view result "
    "summaries, examine growth and risk drivers, visualize timeline data, and compare historical analyses. Unlike "
    "the simpler stateless frontend, the V2 interface introduces session-based access, persistent analysis history, "
    "comparison mode, a modern dashboard layout, and richer user interaction patterns. This report explains the "
    "frontend architecture, component design, state flow, authentication workflow, integration with the FastAPI "
    "backend, and the web technologies used to create a production-style analytical experience. The document is "
    "intended as a subject report that highlights the web engineering contributions of the frontend rather than the "
    "machine learning model internals."
)


KEYWORDS = (
    "React, Vite, Tailwind CSS, Supabase, Recharts, Framer Motion, dashboard UI, authentication, API integration"
)


SECTIONS = [
    (
        "1. Introduction",
        [
            "Frontend V2 of Financial Pragmatic AI was developed as a richer and more interactive web application "
            "layer over the existing financial transcript analysis backend. The primary goal of this frontend is to "
            "transform raw machine learning outputs into a user-friendly analytical dashboard. From an Advanced Web "
            "Technology perspective, the project demonstrates how a modern frontend stack can be used to create an "
            "authenticated single-page application with persistent user history, cloud-backed storage, dynamic "
            "visualization, and reusable component architecture.",
            "The frontend is implemented in React and bundled using Vite, which provides a fast developer workflow "
            "and modern module-based build system. Styling is handled through Tailwind CSS, allowing utility-first "
            "development of a custom dark interface. Additional libraries such as Framer Motion, Lucide React, and "
            "Recharts are used for animation, iconography, and visual analytics. Supabase serves as the client-side "
            "authentication and persistence layer for user sessions and saved analyses.",
            "The significance of this frontend lies in how it combines multiple advanced web technologies into one "
            "cohesive application. The dashboard is not a static page; it reacts to session state, listens to auth "
            "changes, maps backend responses into UI-ready structures, saves analysis results into a cloud database, "
            "and allows side-by-side comparison of historical records. This makes frontend_v2 a strong example of a "
            "modern interactive web application built around real-time state transitions and service integration.",
        ],
    ),
    (
        "2. Objectives of the Frontend V2",
        [
            "The first objective of the V2 frontend is to offer a secure entry layer through email and password "
            "authentication. Users should be able to sign up, log in, and maintain a session without the interface "
            "reloading into separate server-rendered pages. This reflects the single-page application model common "
            "in current frontend systems.",
            "The second objective is to provide a clean financial dashboard through which users can paste a "
            "transcript, upload a local file, run analysis, and view the result using visual and textual summaries. "
            "The dashboard should present signal, score, confidence, distributions, timeline trends, growth drivers, "
            "and risk drivers in a readable format.",
            "The third objective is to introduce history and comparison. Once a user runs an analysis, the record is "
            "stored in Supabase and shown in the sidebar. The user can reopen a previous analysis, update it, or "
            "activate compare mode and select two records for side-by-side analytical comparison.",
            "The final objective is to demonstrate advanced frontend concepts such as reusable components, conditional "
            "rendering, asynchronous API calls, cloud-based session management, animated transitions, responsive "
            "layout design, and client-side file handling.",
        ],
    ),
    (
        "3. Technology Stack Used in Frontend V2",
        [
            "The V2 frontend is based on React 19.2.4 and Vite 8.0.1. React is used for component-driven UI "
            "development and stateful rendering, while Vite handles the development server, ES module workflow, and "
            "production build pipeline.",
            "Tailwind CSS 3.4.13 is used as the styling system. Rather than relying on traditional custom CSS files "
            "for every component, the project uses utility classes to define colors, layout, borders, spacing, "
            "typography, transitions, and responsive behavior directly in JSX. This speeds up UI development and "
            "makes component styling more consistent.",
            "Supabase is used for authentication and persistence. The frontend connects to Supabase using "
            "@supabase/supabase-js and uses environment variables for the project URL and anon key. Framer Motion is "
            "used for smooth content entry and subtle panel transitions. Lucide React provides lightweight icon "
            "components. Recharts is used to render the interactive timeline visualization inside the dashboard.",
        ],
    ),
    (
        "4. Frontend Application Architecture",
        [
            "The V2 application follows a component-based React architecture. The root file App.jsx acts as the "
            "orchestration layer. It manages global page-level state such as transcript text, loading status, current "
            "result, active session, analysis history, comparison selection, active tab, and whether the user is "
            "editing a previously saved analysis. Based on session state, App either renders the Auth component or "
            "the main dashboard shell.",
            "The dashboard layout is divided into three visual regions. The Sidebar component manages recent analyses, "
            "compare mode, record selection, and logout. The Navbar component renders the project branding and active "
            "user information. The main content area contains transcript input, action buttons, the hero summary "
            "panel, compare mode controls, and the tabbed content area.",
            "The tab system is implemented through the Tabs component, which switches between Overview, Insights, "
            "and Compare subviews. Overview displays distributions, timeline chart, and driver panels. Insights is a "
            "simpler panel intended for segment-level textual presentation. Compare renders a detailed comparative "
            "analysis between two selected historical records. This structure shows how a single React application can "
            "be decomposed into focused, reusable, presentation-oriented modules.",
        ],
    ),
    (
        "5. Authentication and Session Handling",
        [
            "Authentication is implemented through Supabase. The supabaseClient.js module creates the client using "
            "the VITE_SUPABASE_URL and VITE_SUPABASE_ANON_KEY environment variables. The Auth component then uses "
            "Supabase methods signInWithPassword and signUp to manage login and registration flows.",
            "The login and sign-up interface is itself a polished frontend module. It uses icons, controlled form "
            "inputs, a loading state, and message state for both error and success conditions. The sign-up flow was "
            "updated so that successful email-confirmation-based registration does not appear as a failure; instead, "
            "the interface explicitly informs the user to check email if a session is not created immediately.",
            "Session persistence is managed in App.jsx using supabase.auth.getSession() and "
            "supabase.auth.onAuthStateChange(). This allows the application to react automatically when a user logs "
            "in, signs out, or restores a previous session. If a session exists, the application loads the user's "
            "analysis history from Supabase. If the session ends, the history list is cleared and the UI returns to "
            "the authentication screen.",
        ],
    ),
    (
        "6. Input Handling and API Integration",
        [
            "The transcript input area is implemented as a controlled textarea bound to React state. This means the "
            "frontend always knows the current transcript value and can reset or update it when the user selects a "
            "previous record from history.",
            "The Analyze action is implemented through an asynchronous fetch request to the FastAPI backend endpoint "
            "at http://localhost:8000/analyze. The frontend sends a JSON payload containing the transcript text. "
            "Once the backend responds, the frontend maps the raw backend structure into a result object that is "
            "better suited for display. For example, intent distribution values are converted into growth, risk, and "
            "neutral percentages, while growth and risk driver arrays are grouped into a drivers object for easier "
            "component consumption.",
            "The file upload feature uses a hidden file input and the browser FileReader API. Instead of uploading "
            "the file directly in frontend_v2, the selected .txt file is read locally in the browser and its content "
            "is inserted into the transcript textarea. This is a useful client-side web technology feature because it "
            "improves user convenience without requiring a separate frontend upload endpoint.",
        ],
    ),
    (
        "7. Cloud-Backed Persistence and History",
        [
            "One of the most important differences between the basic frontend and V2 is history persistence. After a "
            "successful analysis, the result is stored in a Supabase table named analyses. The stored record includes "
            "the user identifier, transcript, signal, score, distribution, growth drivers, risk drivers, and "
            "timeline.",
            "When App detects a valid session, it calls fetchHistory and retrieves the user's saved analyses from "
            "Supabase, ordered by creation date. These records are displayed in the Sidebar component. Each record "
            "shows a compact signal label, a timestamp, and a short transcript preview. Clicking a history item "
            "restores the transcript and its saved result back into the main dashboard. This demonstrates practical "
            "use of frontend-to-database interaction in a real web application.",
            "The application also supports editing existing history entries. If a user opens an old record and runs a "
            "new analysis from that context, the application can update the existing Supabase row instead of creating "
            "a fresh record. This is controlled through selectedAnalysis and isFromHistory state variables, showing a "
            "useful example of CRUD-oriented web behavior inside a dashboard UI.",
        ],
    ),
    (
        "8. Comparison Mode and Advanced Interaction Design",
        [
            "Compare mode is one of the most advanced frontend features in V2. When the user enables compare mode, "
            "the Sidebar changes from simple history navigation into a two-slot selection interface. The first chosen "
            "record becomes A and the second becomes B. If the user keeps selecting more records, the compare window "
            "slides forward by replacing older selections.",
            "The Compare component then renders a structured side-by-side analysis. It computes score deltas, "
            "distribution changes, interpretation labels, and driver transitions. A dedicated helper inside the "
            "component compares previous and current drivers and labels them as NEW, REMOVED, or retained. This is a "
            "good example of how frontend components can implement meaningful interpretation logic instead of only "
            "displaying raw backend fields.",
            "Timeline-level insight is also included in the compare view. The component derives a stability or "
            "volatility summary by examining changes in the stored timeline. This reinforces the Advanced Web "
            "Technology theme of building smart, interactive interfaces that derive added value from data on the "
            "client side.",
        ],
    ),
    (
        "9. Data Visualization and User Interface Design",
        [
            "The visual design of frontend_v2 follows a dark fintech dashboard style. Tailwind CSS is used to build "
            "glassmorphism-inspired cards with translucent backgrounds, border definition, blur effects, and smooth "
            "transitions. The hero panel changes glow color based on signal type, using green for growth, red for "
            "risk, and blue for neutral. This creates immediate visual feedback.",
            "The Overview component presents multiple layers of visualization. Distribution bars show the percentage "
            "mix of growth, risk, and neutral interpretations. TimelineChart uses Recharts AreaChart to render the "
            "trajectory of the transcript across segments, with zone overlays, custom dots, custom tooltips, and a "
            "trend badge. This is a strong demonstration of how charting libraries can be integrated into React "
            "components for dynamic analytical dashboards.",
            "Animation is used carefully rather than excessively. Framer Motion is applied to fade in content, while "
            "buttons and panels use CSS transitions and hover states. Icons from Lucide React improve navigational "
            "clarity without adding visual clutter. Overall, the frontend reflects deliberate UI engineering rather "
            "than minimal form-based web design.",
        ],
    ),
    (
        "10. Responsive and Interactive Frontend Behavior",
        [
            "The project demonstrates multiple advanced frontend interaction patterns. Controlled inputs are used for "
            "forms and transcript editing. Conditional rendering is used to switch between Auth, Overview, Compare, "
            "and other states. Loading state changes button labels such as Analyzing... to reflect asynchronous "
            "operations clearly to the user.",
            "The application also includes hidden file input triggering through a custom button, real-time session "
            "state transitions, interactive history selection, toggle-based compare mode, and context-sensitive "
            "hero-panel updates. These behaviors are implemented through standard React state hooks and event "
            "handlers, showing a practical application of event-driven frontend design.",
            "The layout itself is split into sidebar and content panel regions, and most cards use flexible width and "
            "responsive spacing. Although the main experience is optimized for desktop dashboard use, the use of "
            "utility classes and flex/grid layout makes the application adaptable and maintainable.",
        ],
    ),
    (
        "11. Challenges Faced in Frontend V2",
        [
            "A major challenge in the V2 frontend was aligning backend response structure with frontend display "
            "requirements. The backend returns detailed analysis information, but the frontend expects grouped "
            "distributions, driver categories, and timeline data in component-friendly shapes. This required response "
            "mapping inside App.jsx before state could be updated safely.",
            "Another challenge was handling Supabase authentication and persistence smoothly. Signup behavior, session "
            "restoration, logout transitions, and history refresh logic all had to be coordinated carefully. Without "
            "clear success and error messages, users could misinterpret expected confirmation flows as failures.",
            "A third challenge was designing a visually strong dashboard while keeping the component structure "
            "manageable. This was solved by dividing the app into small reusable modules such as Sidebar, Navbar, "
            "Tabs, Overview, Compare, and TimelineChart, each with a focused responsibility.",
        ],
    ),
    (
        "12. Conclusion",
        [
            "Frontend V2 of Financial Pragmatic AI is a strong example of Advanced Web Technology applied to a real "
            "analytical product. It uses a modern frontend stack consisting of React, Vite, Tailwind CSS, Supabase, "
            "Framer Motion, Lucide React, and Recharts to deliver an authenticated, cloud-connected, interactive "
            "dashboard experience.",
            "From a subject perspective, the project demonstrates component-based design, client-side authentication, "
            "state-driven rendering, asynchronous API integration, file handling in the browser, cloud persistence, "
            "data visualization, and UI responsiveness. These are all important concepts in modern web application "
            "development.",
            "The V2 frontend goes beyond simple page rendering and shows how advanced frontend engineering can make a "
            "machine-learning-backed system usable, understandable, and visually effective for end users. For this "
            "reason, it is an appropriate and well-scoped project to discuss in an Advanced Web Technology report.",
        ],
    ),
]


STACK_ROWS = [
    ("React 19.2.4", "Component-based UI development and state-driven rendering"),
    ("Vite 8.0.1", "Fast development server and modern frontend build pipeline"),
    ("Tailwind CSS 3.4.13", "Utility-first styling and dashboard layout design"),
    ("Supabase JS 2.101.1", "Authentication, session handling, and history persistence"),
    ("Framer Motion", "Animated content transitions"),
    ("Lucide React", "Icon library for UI clarity"),
    ("Recharts 3.8.1", "Interactive timeline and analytical chart visualization"),
]


COMPONENT_ROWS = [
    ("App.jsx", "Global state management, auth gating, backend calls, and persistence orchestration"),
    ("Auth.jsx", "Login and sign-up interface with session messages"),
    ("Sidebar.jsx", "History browsing, compare mode selection, and logout"),
    ("Navbar.jsx", "Top dashboard identity and user session display"),
    ("Tabs.jsx", "Navigation between overview, insights, and compare views"),
    ("Overview.jsx", "Distribution, drivers, and timeline display"),
    ("Compare.jsx", "Record comparison logic and delta rendering"),
    ("TimelineChart.jsx", "Interactive timeline visualization using Recharts"),
    ("supabaseClient.js", "Supabase connection bootstrap"),
]


REFERENCES = [
    "1. React Documentation, https://react.dev/",
    "2. Vite Documentation, https://vite.dev/",
    "3. Tailwind CSS Documentation, https://tailwindcss.com/docs",
    "4. Supabase Documentation, https://supabase.com/docs",
    "5. Recharts Documentation, https://recharts.org/",
    "6. Framer Motion Documentation, https://www.framer.com/motion/",
    "7. Lucide Documentation, https://lucide.dev/",
]


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_paragraph(doc: Document, text: str, style: str = "IEEE Paragraph") -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_table(doc: Document, rows: list[tuple[str, str]], heading: str) -> None:
    caption = doc.add_paragraph(style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(heading)

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Purpose"

    for paragraph in table.cell(0, 0).paragraphs + table.cell(0, 1).paragraphs:
        for run in paragraph.runs:
            run.bold = True

    for index, (item, purpose) in enumerate(rows, start=1):
        table.cell(index, 0).text = item
        table.cell(index, 1).text = purpose

    doc.add_paragraph("", style="Normal")


def build_report() -> Path:
    doc = Document(str(TEMPLATE_PATH))
    clear_document(doc)

    title = doc.add_paragraph(style="paper title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    author = doc.add_paragraph(style="Author")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run(AUTHOR)

    for line in AFFILIATION_LINES:
        paragraph = doc.add_paragraph(style="Affiliation")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)

    sep = doc.add_paragraph(style="Affiliation")
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.add_run("_" * 110)

    abstract = doc.add_paragraph(style="Normal")
    abstract.add_run("Abstract - ").bold = True
    abstract.add_run(ABSTRACT)

    doc.add_paragraph("", style="Normal")

    keywords = doc.add_paragraph(style="Normal")
    keywords.add_run("Index Terms - ").bold = True
    keywords.add_run(KEYWORDS)

    sep2 = doc.add_paragraph(style="Normal")
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep2.add_run("_" * 110)
    doc.add_paragraph("", style="Normal")

    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)

        if heading == "3. Technology Stack Used in Frontend V2":
            add_table(doc, STACK_ROWS, "Table 1. Frontend V2 Technology Stack")

        if heading == "4. Frontend Application Architecture":
            add_table(doc, COMPONENT_ROWS, "Table 2. Main Frontend V2 Components")

    add_heading(doc, "13. References")
    for reference in REFERENCES:
        doc.add_paragraph(reference, style="Normal")

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_report()
    print(output)
